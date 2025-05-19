from docx import Document
from docx.shared import Pt
import io

def create_transparency_report(company_name, url_context_raw, url_context_sum, 
                               additional_context_raw, additional_context_sum, 
                               lead_magnet_raw, lead_magnet_sum):
    doc = Document()
    
    # Title
    title = doc.add_heading(f"{company_name} AI Transparency Report", level=1)
    title.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(f"This document outlines the source text and AI-generated summaries used for creating ad copy for {company_name}.")
    doc.add_paragraph() # Spacer

    # Helper to add sections
    def add_section(heading_text, raw_text, sum_text):
        doc.add_heading(heading_text, level=2)
        p_raw = doc.add_paragraph()
        p_raw.add_run(raw_text if raw_text else "Not provided.").font.size = Pt(10)
        
        doc.add_heading("AI Summarized:", level=3)
        p_sum = doc.add_paragraph()
        p_sum.add_run(sum_text if sum_text else ("Not provided." if not raw_text else "Summary not generated.")).font.size = Pt(10)
        doc.add_paragraph()

    add_section("Raw Extract from URL:", url_context_raw, url_context_sum)
    add_section("Extract from Additional Upload:", additional_context_raw, additional_context_sum)
    add_section("Extract from Lead Magnet:", lead_magnet_raw, lead_magnet_sum)

    # Save to a BytesIO object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream